"""파일 포맷별 텍스트 추출 - 검출 엔진에 통일된 텍스트를 공급.

통일 인터페이스 `extract_text(path) -> str` 로 모든 포맷을 텍스트로 받는다.
포맷별 라이브러리를 지연 import 하여, 라이브러리가 없거나 파싱이 실패해도
ExtractionError 로 격리되어 전체 스캔이 멈추지 않는다(기획서 9장 리스크 #3).

기본 동작(외부 의존성 없이):
    - 평문/소스/로그: 표준 라이브러리(인코딩 자동 감지)
    - hwpx: zip+xml 파싱(표준 라이브러리)  ← 국내 SI 차별점
    - xlsx: openpyxl 있으면 사용, 없으면 zip+xml 폴백

선택 라이브러리(있으면 활성):
    openpyxl, xlrd, python-docx, pdfplumber, olefile, pytesseract, Pillow, chardet
"""

from __future__ import annotations

import csv
import io
import json
import re
import struct
import zipfile
import zlib
from dataclasses import dataclass, field as _dc_field
from pathlib import Path

__all__ = ["extract_text", "extract_doc", "ExtractedDoc", "ExtractionError", "is_supported"]


class ExtractionError(Exception):
    """파싱 불가(암호/손상/미지원/라이브러리 없음) 시 발생 → '검사불가' 처리."""


@dataclass
class ExtractedDoc:
    """추출 결과. 검출 엔진이 그대로 스캔하는 평문(text)과, 구조화 포맷에서
    얻은 필드 구간(fields=[(시작, 끝, 라벨)])을 함께 제공한다.

    fields 의 라벨(컬럼 헤더·JSON 키)을 통해 엔진이 '각 확장자 구조'에 맞춘
    문맥 검증을 수행한다(예: '주민등록번호' 열의 값은 강하게 그 유형으로 추정)."""

    text: str
    fields: list[tuple[int, int, str]] = _dc_field(default_factory=list)


class _DocBuilder:
    """평문과 필드 구간을 함께 누적하는 보조 빌더(오프셋 정합 보장)."""

    def __init__(self) -> None:
        self._parts: list[str] = []
        self._pos = 0
        self.fields: list[tuple[int, int, str]] = []

    def add(self, value, label: str = "") -> None:
        """값 한 토큰을 한 줄로 추가하고, 라벨이 있으면 그 구간을 기록한다."""
        if value is None:
            return
        s = str(value).strip()
        if not s:
            return
        start = self._pos
        self._parts.append(s)
        self._pos += len(s)
        if label:
            self.fields.append((start, self._pos, label.strip()))
        self._parts.append("\n")
        self._pos += 1

    def build(self) -> ExtractedDoc:
        return ExtractedDoc("".join(self._parts), self.fields)


SUPPORTED_TEXT = {
    # 문서/로그/설정
    ".txt", ".csv", ".tsv", ".log", ".json", ".jsonl", ".ndjson", ".xml",
    ".md", ".yml", ".yaml", ".ini", ".cfg", ".conf", ".toml",
    ".properties", ".env", ".html", ".htm",
    # 소스코드(SI 개발자 산출물)
    ".py", ".java", ".js", ".ts", ".tsx", ".jsx", ".kt", ".go", ".c", ".cpp",
    ".h", ".hpp", ".cs", ".php", ".rb", ".rs", ".swift", ".scala", ".pl", ".r",
    ".sh", ".bat", ".ps1",
    # DB/백업 — 텍스트 덤프(SQL 스크립트, mysqldump/pg_dump 등)
    ".sql", ".ddl", ".dump", ".bak",
}
SUPPORTED_OFFICE = {".xlsx", ".xls", ".docx"}
SUPPORTED_HWP = {".hwp", ".hwpx"}
SUPPORTED_PDF = {".pdf"}
SUPPORTED_IMAGE = {".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".tif"}
# DB 파일(백업된 고객 데이터) — SQLite 바이너리. SI 개발자 PC 대응.
SUPPORTED_DB = {".db", ".sqlite", ".sqlite3", ".db3"}
# 디자인 파일(디자이너 직무). 실제 추출은 design_extractors 가 담당.
SUPPORTED_DESIGN = {".psd", ".psb", ".xd"}

_ALL_SUPPORTED = (
    SUPPORTED_TEXT | SUPPORTED_OFFICE | SUPPORTED_HWP | SUPPORTED_PDF
    | SUPPORTED_IMAGE | SUPPORTED_DB | SUPPORTED_DESIGN
)


def is_supported(path: str | Path) -> bool:
    return Path(path).suffix.lower() in _ALL_SUPPORTED


#: 추출 시 메모리에 통째로 올리는 포맷의 파일 크기 상한(초과 시 검사불가).
#: SQLite(.db 등)는 행 상한으로 스트리밍하므로 이 제한에서 제외한다.
MAX_EXTRACT_BYTES = 200 * 1024 * 1024  # 200MB


def _guard_size(path: Path) -> None:
    if path.suffix.lower() in SUPPORTED_DB:
        return
    try:
        size = path.stat().st_size
    except OSError:
        return
    if size > MAX_EXTRACT_BYTES:
        raise ExtractionError(
            f"파일이 너무 커서 건너뜀 "
            f"({size // (1024 * 1024)}MB > {MAX_EXTRACT_BYTES // (1024 * 1024)}MB)")


def extract_text(path: str | Path, ocr_enabled: bool = True) -> str:
    """확장자에 맞는 추출기로 분기. 실패 시 ExtractionError."""
    path = Path(path)
    ext = path.suffix.lower()
    try:
        _guard_size(path)
        if ext in SUPPORTED_TEXT:
            return _extract_plain(path)
        if ext == ".xlsx":
            return _extract_xlsx(path)
        if ext == ".xls":
            return _extract_xls(path)
        if ext == ".docx":
            return _extract_docx(path)
        if ext == ".hwpx":
            return _extract_hwpx(path)
        if ext == ".hwp":
            return _extract_hwp_ole(path)
        if ext == ".pdf":
            return _extract_pdf(path, ocr_enabled)
        if ext in SUPPORTED_DB:
            return _extract_sqlite(path)
        if ext in SUPPORTED_IMAGE:
            return _extract_image(path) if ocr_enabled else ""
        if ext in SUPPORTED_DESIGN:
            from .design_extractors import extract_design_text

            return extract_design_text(path, ocr_enabled)
        raise ExtractionError(f"미지원 형식: {ext}")
    except ExtractionError:
        raise
    except Exception as e:  # 포맷 라이브러리의 모든 예외를 격리
        raise ExtractionError(f"{path.name} 파싱 실패: {e}") from e


def extract_doc(path: str | Path, ocr_enabled: bool = True) -> ExtractedDoc:
    """구조 인식 추출. 표(csv/xlsx/xls)·JSON 은 필드 라벨까지 함께 돌려준다.

    구조화 파싱이 실패하면 평문 추출로 안전하게 폴백한다(필드 라벨 없음).
    그 외 포맷은 기존 extract_text 결과를 라벨 없는 ExtractedDoc 로 감싼다."""
    path = Path(path)
    ext = path.suffix.lower()
    _guard_size(path)  # 초과 시 ExtractionError → 검사불가
    try:
        if ext == ".csv":
            return _doc_csv(path)
        if ext == ".tsv":
            return _doc_csv(path, delimiter="\t")
        if ext == ".json":
            return _doc_json(path)
        if ext == ".xlsx":
            return _doc_xlsx(path)
        if ext == ".xls":
            return _doc_xls(path)
        if ext in SUPPORTED_DB:
            return _doc_sqlite(path)
    except ExtractionError:
        raise
    except Exception:
        # 구조 파싱 실패 → 평문으로 폴백(검출 자체는 계속 수행)
        pass
    return ExtractedDoc(extract_text(path, ocr_enabled=ocr_enabled), [])


# ---------------------------------------------------------------------------
# 구조 인식 추출 (표/JSON) — 필드 라벨 부착
# ---------------------------------------------------------------------------
def _doc_csv(path: Path, delimiter: str = ",") -> ExtractedDoc:
    """CSV/TSV: 첫 비어있지 않은 행을 헤더로 삼아 각 셀에 컬럼 라벨을 부여."""
    raw = path.read_bytes()
    if not raw:
        return ExtractedDoc("", [])
    text = _decode_bytes(raw)
    reader = csv.reader(io.StringIO(text), delimiter=delimiter)
    rows = [r for r in reader]
    header: list[str] = []
    b = _DocBuilder()
    for row in rows:
        if not any(c.strip() for c in row):
            continue
        if not header:
            header = [c.strip() for c in row]
            continue
        for i, cell in enumerate(row):
            label = header[i] if i < len(header) else ""
            b.add(cell, label)
    return b.build()


def _doc_json(path: Path) -> ExtractedDoc:
    """JSON: 모든 말단 값에 그 키(객체) 또는 부모 키(배열)를 라벨로 부여."""
    raw = path.read_bytes()
    data = json.loads(_decode_bytes(raw)) if raw else {}
    b = _DocBuilder()

    def walk(node, label: str) -> None:
        if isinstance(node, dict):
            for k, v in node.items():
                walk(v, str(k))
        elif isinstance(node, list):
            for item in node:
                walk(item, label)
        elif isinstance(node, bool) or node is None:
            return
        else:
            b.add(node, label)

    walk(data, "")
    return b.build()


def _doc_xlsx(path: Path) -> ExtractedDoc:
    """XLSX: openpyxl 있으면 시트별 첫 행을 헤더로 구조 추출, 없으면 평문 폴백."""
    try:
        import openpyxl  # type: ignore
    except ImportError:
        return ExtractedDoc(_extract_xlsx_stdlib(path), [])

    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    b = _DocBuilder()
    try:
        for ws in wb.worksheets:
            header: list[str] = []
            for row in ws.iter_rows(values_only=True):
                cells = list(row)
                if not any(c is not None and str(c).strip() for c in cells):
                    continue
                if not header:
                    header = [("" if c is None else str(c).strip()) for c in cells]
                    continue
                for i, c in enumerate(cells):
                    label = header[i] if i < len(header) else ""
                    b.add(c, label)
    finally:
        wb.close()
    return b.build()


def _doc_xls(path: Path) -> ExtractedDoc:
    """구형 XLS: xlrd 로 시트별 첫 행을 헤더로 구조 추출. 없으면 ExtractionError."""
    try:
        import xlrd  # type: ignore
    except ImportError:
        raise ExtractionError("xls 파싱에 xlrd 라이브러리가 필요합니다")
    book = xlrd.open_workbook(path)
    b = _DocBuilder()
    for sheet in book.sheets():
        header: list[str] = []
        for r in range(sheet.nrows):
            cells = sheet.row_values(r)
            if not any(str(c).strip() for c in cells):
                continue
            if not header:
                header = [str(c).strip() for c in cells]
                continue
            for i, c in enumerate(cells):
                label = header[i] if i < len(header) else ""
                b.add(c, label)
    return b.build()


# ---------------------------------------------------------------------------
# DB 파일(백업된 고객 데이터) — SQLite. SI 개발자 PC의 .db/.sqlite 대응.
# ---------------------------------------------------------------------------
_SQLITE_MAGIC = b"SQLite format 3\x00"
_DB_ROW_LIMIT = 50000           # 테이블당 최대 행(과도한 메모리 방지)


def _is_sqlite(path: Path) -> bool:
    try:
        with open(path, "rb") as f:
            return f.read(16) == _SQLITE_MAGIC
    except OSError:
        return False


def _sqlite_tables(cur) -> list[str]:
    return [r[0] for r in cur.execute(
        "SELECT name FROM sqlite_master WHERE type='table' "
        "AND name NOT LIKE 'sqlite_%'")]


def _extract_sqlite(path: Path) -> str:
    """SQLite DB의 모든 테이블 텍스트 셀을 추출. SQLite가 아니면 평문 폴백."""
    if not _is_sqlite(path):
        return _extract_plain(path)   # .db 가 텍스트일 수도 있음
    import sqlite3
    parts: list[str] = []
    conn = sqlite3.connect(f"file:{path}?mode=ro", uri=True)
    try:
        cur = conn.cursor()
        for t in _sqlite_tables(cur):
            try:
                rows = cur.execute(
                    f'SELECT * FROM "{t}" LIMIT {_DB_ROW_LIMIT}').fetchall()
            except sqlite3.Error:
                continue
            for row in rows:
                cells = [str(c) for c in row
                         if c is not None and not isinstance(c, (bytes, bytearray))]
                if cells:
                    parts.append(" ".join(cells))
    finally:
        conn.close()
    return "\n".join(parts)


def _doc_sqlite(path: Path) -> ExtractedDoc:
    """SQLite DB를 컬럼명 라벨과 함께 구조 추출(예: 'rrn'/'phone' 컬럼 → 필드 구제)."""
    if not _is_sqlite(path):
        return ExtractedDoc(_extract_plain(path), [])
    import sqlite3
    b = _DocBuilder()
    conn = sqlite3.connect(f"file:{path}?mode=ro", uri=True)
    try:
        cur = conn.cursor()
        for t in _sqlite_tables(cur):
            try:
                cur.execute(f'SELECT * FROM "{t}" LIMIT {_DB_ROW_LIMIT}')
            except sqlite3.Error:
                continue
            cols = [d[0] for d in (cur.description or [])]
            for row in cur.fetchall():
                for i, c in enumerate(row):
                    if c is None or isinstance(c, (bytes, bytearray)):
                        continue
                    b.add(c, cols[i] if i < len(cols) else "")
    finally:
        conn.close()
    return b.build()


# ---------------------------------------------------------------------------
# 평문 / 소스 / 로그
# ---------------------------------------------------------------------------
def _extract_plain(path: Path) -> str:
    """텍스트/소스/로그 - 인코딩 자동 감지(한글 cp949 대응)."""
    raw = path.read_bytes()
    if not raw:
        return ""
    return _decode_bytes(raw)


def _decode_bytes(raw: bytes) -> str:
    # chardet 있으면 1차로 활용
    try:
        import chardet  # type: ignore

        guessed = chardet.detect(raw[:100_000]).get("encoding")
        if guessed:
            try:
                return raw.decode(guessed, errors="replace")
            except LookupError:
                pass
    except ImportError:
        pass
    # 폴백: 흔한 인코딩 순차 시도(국내 SI는 cp949/euc-kr 빈번)
    for enc in ("utf-8-sig", "utf-8", "cp949", "euc-kr"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


# ---------------------------------------------------------------------------
# 엑셀
# ---------------------------------------------------------------------------
def _extract_xlsx(path: Path) -> str:
    """엑셀(xlsx). openpyxl 있으면 사용, 없으면 zip+xml 폴백."""
    try:
        import openpyxl  # type: ignore
    except ImportError:
        return _extract_xlsx_stdlib(path)

    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    parts: list[str] = []
    try:
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None]
                if cells:
                    parts.append(" ".join(cells))
    finally:
        wb.close()
    return "\n".join(parts)


def _extract_xlsx_stdlib(path: Path) -> str:
    """openpyxl 없이 xlsx에서 텍스트 추출(검출 목적의 폴백).

    sharedStrings 의 <t> 텍스트와 각 시트의 <v> 값을 모아 검출 엔진에 넘긴다.
    숫자로 저장된 PII(예: 숫자형 셀의 카드번호)도 <v> 로 포착된다.
    """
    if not zipfile.is_zipfile(path):
        raise ExtractionError("유효한 xlsx(zip) 가 아님")
    parts: list[str] = []
    with zipfile.ZipFile(path) as z:
        names = z.namelist()
        if "xl/sharedStrings.xml" in names:
            xml = z.read("xl/sharedStrings.xml").decode("utf-8", errors="replace")
            parts += re.findall(r"<t[^>]*>(.*?)</t>", xml, re.DOTALL)
        for name in names:
            if name.startswith("xl/worksheets/") and name.endswith(".xml"):
                xml = z.read(name).decode("utf-8", errors="replace")
                parts += re.findall(r"<v>(.*?)</v>", xml, re.DOTALL)
    return "\n".join(_strip_xml(p) for p in parts)


def _extract_xls(path: Path) -> str:
    """구형 엑셀(xls) - xlrd 필요."""
    try:
        import xlrd  # type: ignore
    except ImportError:
        raise ExtractionError("xls 파싱에 xlrd 라이브러리가 필요합니다")
    book = xlrd.open_workbook(path)
    parts: list[str] = []
    for sheet in book.sheets():
        for r in range(sheet.nrows):
            cells = [str(c) for c in sheet.row_values(r) if c != ""]
            if cells:
                parts.append(" ".join(cells))
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# 워드
# ---------------------------------------------------------------------------
def _extract_docx(path: Path) -> str:
    """워드(docx) - 본문 단락 + 표 셀. python-docx 필요."""
    try:
        import docx  # type: ignore
    except ImportError:
        raise ExtractionError("docx 파싱에 python-docx 라이브러리가 필요합니다")
    doc = docx.Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" ".join(cells))
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# 한글(hwp/hwpx) - 국내 SI 환경 필수
# ---------------------------------------------------------------------------
def _extract_hwpx(path: Path) -> str:
    """hwpx: zip 내부 Contents/section*.xml 의 <hp:t> 텍스트 추출(표준 라이브러리)."""
    if not zipfile.is_zipfile(path):
        raise ExtractionError("유효한 hwpx(zip) 가 아님")
    parts: list[str] = []
    with zipfile.ZipFile(path) as z:
        for name in z.namelist():
            if "section" in name.lower() and name.endswith(".xml"):
                xml = z.read(name).decode("utf-8", errors="replace")
                parts += re.findall(r"<hp:t>(.*?)</hp:t>", xml, re.DOTALL)
    return "\n".join(_strip_xml(p) for p in parts)


def _extract_hwp_ole(path: Path) -> str:
    """구형 hwp: OLE BodyText 스트림(zlib 압축) 해제 후 텍스트 추출. olefile 필요."""
    try:
        import olefile  # type: ignore
    except ImportError:
        raise ExtractionError("hwp 파싱에 olefile 라이브러리가 필요합니다")

    if not olefile.isOleFile(str(path)):
        raise ExtractionError("유효한 HWP(OLE) 파일이 아님")
    ole = olefile.OleFileIO(str(path))
    try:
        header = ole.openstream("FileHeader").read()
        is_compressed = bool(header[36] & 0x01)
        is_encrypted = bool(header[36] & 0x02)
        if is_encrypted:
            raise ExtractionError("암호로 보호된 HWP 문서")

        texts: list[str] = []
        for entry in ole.listdir():
            if entry[0] == "BodyText":
                data = ole.openstream(entry).read()
                if is_compressed:
                    data = zlib.decompress(data, -15)
                texts.append(_decode_hwp_records(data))
        return "\n".join(texts)
    finally:
        ole.close()


def _decode_hwp_records(data: bytes) -> str:
    """HWP 레코드 구조에서 PARA_TEXT(태그 67, UTF-16LE) 추출."""
    out: list[str] = []
    i, n = 0, len(data)
    while i + 4 <= n:
        header = struct.unpack("<I", data[i : i + 4])[0]
        tag = header & 0x3FF
        size = (header >> 20) & 0xFFF
        i += 4
        if tag == 67:  # HWPTAG_PARA_TEXT
            chunk = data[i : i + size]
            out.append(chunk.decode("utf-16-le", errors="replace"))
        i += size
    return "".join(out)


# ---------------------------------------------------------------------------
# PDF / 이미지(OCR)
# ---------------------------------------------------------------------------
def _extract_pdf(path: Path, ocr_enabled: bool) -> str:
    """PDF - 텍스트 레이어 우선, 없으면(스캔 PDF) OCR 폴백. pdfplumber 필요."""
    try:
        import pdfplumber  # type: ignore
    except ImportError:
        raise ExtractionError("pdf 파싱에 pdfplumber 라이브러리가 필요합니다")

    parts: list[str] = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            parts.append(page.extract_text() or "")
    text = "\n".join(parts).strip()
    if not text and ocr_enabled:
        return _ocr_pdf(path)
    return text


_TESSERACT_PATHS = [
    r"C:\Program Files\Tesseract-OCR\tesseract.exe",
    r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
    "/usr/bin/tesseract",
]
_ocr_ready = False
_ocr_lang = "eng"


def _configure_ocr() -> str:
    """Tesseract 실행 경로·언어 자동 설정. 사용 가능 언어(kor 포함 여부) 반환."""
    global _ocr_ready, _ocr_lang
    if _ocr_ready:
        return _ocr_lang
    import os

    import pytesseract  # type: ignore

    for p in _TESSERACT_PATHS:
        if os.path.exists(p):
            pytesseract.pytesseract.tesseract_cmd = p
            break
    # kor 포함 사용자 tessdata 우선
    user_td = os.path.join(
        os.environ.get("LOCALAPPDATA", ""), "SoliGuard", "tessdata")
    if os.path.isdir(user_td):
        os.environ["TESSDATA_PREFIX"] = user_td
    try:
        langs = set(pytesseract.get_languages(config=""))
    except Exception:
        langs = set()
    _ocr_lang = "kor+eng" if "kor" in langs else "eng"
    _ocr_ready = True
    return _ocr_lang


def _extract_image(path: Path) -> str:
    """이미지 OCR - 신분증·계약서 스캔본(로컬 Tesseract). pytesseract 필요."""
    try:
        import pytesseract  # type: ignore
        from PIL import Image  # type: ignore
    except ImportError:
        raise ExtractionError("이미지 OCR에 pytesseract/Pillow 가 필요합니다")
    lang = _configure_ocr()
    img = Image.open(str(path))
    return pytesseract.image_to_string(img, lang=lang)


#: 스캔 PDF OCR 최대 페이지(대용량 스캔본의 시간·메모리 폭주 방지)
MAX_OCR_PAGES = 30


def _ocr_pdf(path: Path) -> str:
    """스캔 PDF를 이미지로 변환 후 OCR. pdf2image + pytesseract 필요.

    선두 MAX_OCR_PAGES 페이지까지만 처리한다(대용량 스캔본 보호)."""
    try:
        from pdf2image import convert_from_path  # type: ignore
        import pytesseract  # type: ignore
    except ImportError:
        raise ExtractionError("스캔 PDF OCR에 pdf2image/pytesseract 가 필요합니다")
    lang = _configure_ocr()
    images = convert_from_path(
        str(path), dpi=200, first_page=1, last_page=MAX_OCR_PAGES)
    parts = [pytesseract.image_to_string(img, lang=lang) for img in images]
    if len(images) >= MAX_OCR_PAGES:
        parts.append(f"[안내] OCR은 선두 {MAX_OCR_PAGES}페이지까지만 검사했습니다.")
    return "\n".join(parts)


# ---------------------------------------------------------------------------
def _strip_xml(text: str) -> str:
    """잔여 XML 태그 제거 + 기본 엔티티 복원."""
    text = re.sub(r"<[^>]+>", "", text)
    return (
        text.replace("&amp;", "&")
        .replace("&lt;", "<")
        .replace("&gt;", ">")
        .replace("&quot;", '"')
        .replace("&#xD;", "")
    )
